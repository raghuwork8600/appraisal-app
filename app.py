# import streamlit as st
# import pandas as pd
# from docx import Document
# import os
# from io import BytesIO
# import pythoncom
# import zipfile
# import shutil
# import win32com.client
# import time

# st.title(":blue[Appraisal Letter :orange[Generator]")

# # Function to convert DOCX to PDF using win32com
# def convert_docx_to_pdf(docx_file, pdf_path):
#     pythoncom.CoInitialize()
#     word = None
#     doc = None
#     try:
#         word = win32com.client.Dispatch("Word.Application")
#         word.Visible = False
#         doc = word.Documents.Open(docx_file)
#         doc.SaveAs(pdf_path, FileFormat=17) # 17 is for PDF
#     except Exception as e:
#         print(f"Error converting {docx_file} to PDF: {e}")
#     finally:
#         if doc:
#             doc.Close(False)
#         if word:
#             word.Quit()
#         pythoncom.CoUninitialize()
        
        

# # Function to replace placeholders in the document
# def replace_placeholders_in_doc(doc, placeholders):
#     # Replace in paragraphs
#     for paragraph in doc.paragraphs:
#         for placeholder, value in placeholders.items():
#             if placeholder in paragraph.text:
#                 paragraph.text = paragraph.text.replace(placeholder, str(value))
    
#     # Replace in table cells
#     for table in doc.tables:
#         rows_to_remove = []
#         for row_index, row in enumerate(table.rows):
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     for placeholder, value in placeholders.items():
#                         if placeholder in paragraph.text:
#                             paragraph.text = paragraph.text.replace(placeholder, str(value))
#                             if placeholder == '<<RA>>' and value == '':
#                                 rows_to_remove.append(row_index)
        
#         for row_index in reversed(rows_to_remove):
#             tbl = table._tbl
#             tbl.remove(tbl.tr_lst[row_index])

# # Function to format numbers in Indian numbering system
# def format_number_indian(number):
#     # Convert the number to a string
#     num_str = str(number)
    
#     # Split the number into integer and decimal parts
#     if '.' in num_str:
#         integer_part, decimal_part = num_str.split('.')
#     else:
#         integer_part, decimal_part = num_str, ''
    
#     # Reverse the integer part for easier processing
#     integer_part = integer_part[::-1]
    
#     # Insert commas according to the Indian numbering system
#     groups = []
#     groups.append(integer_part[:3])
#     integer_part = integer_part[3:]
    
#     while integer_part:
#         groups.append(integer_part[:2])
#         integer_part = integer_part[2:]
    
#     # Reverse the groups and join them with commas
#     formatted_integer = ','.join(groups)[::-1]
    
#     # Combine the integer and decimal parts
#     if decimal_part:
#         return formatted_integer + '.' + decimal_part
#     else:
#         return formatted_integer


# # Function to zip files
# def zip_files(file_paths, zip_path):
#     with zipfile.ZipFile(zip_path, 'w') as zipf:
#         for file in file_paths:
#             zipf.write(file, os.path.basename(file))

# # Function to clean up temporary files
# def cleanup_temp_files(temp_dir):
#     if os.path.exists(temp_dir):
#         # Remove all files in the directory
#         for file in os.listdir(temp_dir):
#             file_path = os.path.join(temp_dir, file)
#             if os.path.isfile(file_path):
#                 os.remove(file_path)
        
#         # Remove the directory itself
#         os.rmdir(temp_dir)

# # Initialize session state
# if 'doc_paths' not in st.session_state:
#     st.session_state.doc_paths = []
# if 'pdf_paths' not in st.session_state:
#     st.session_state.pdf_paths = []
# if 'doc_zip_path' not in st.session_state:
#     st.session_state.doc_zip_path = ""
# if 'pdf_zip_path' not in st.session_state:
#     st.session_state.pdf_zip_path = ""
# if 'cleanup' not in st.session_state:
#     st.session_state.cleanup = False

# # Upload the Excel file
# uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")

# if uploaded_file and not st.session_state.doc_paths:
#     # Define and create temporary directory
#     temp_dir = "temp_files"
#     os.makedirs(temp_dir, exist_ok=True)

#     doc_paths = []
#     pdf_paths = []
#     start_time = time.time()
#     print(f"Start Time: {start_time}")


#     with st.spinner('Processing the Excel file...'):
#         df = pd.read_excel(uploaded_file)
#         template_path = 'Compensation Revision Letter_Format.docx'
#         excel_read_time = time.time()
#         print(f"Time to read excel: {excel_read_time - start_time}")

#         for index, row in df.iterrows():
#             doc = Document(template_path)

#             fixed = int(str(row['Fixed Pay']).replace(',', ''))
#             variable = int(str(row['Variable Pay']).replace(',', ''))
#             retention = 0 if row['Retention Pay'] == 'No' else int(str(row['Retention Pay']).replace(',', ''))
#             total_compensation = fixed + variable + retention

#             placeholders = {
#                 '<<Month YYYY>>': pd.to_datetime(row['Month of the Letter issued']).strftime("%d %b %Y"),
#                 '<<Name>>': row['Name of Employee'],
#                 '<<Designation>>': row['New Designation'] if row['Promotion'] == 'Yes' else '',
#                 '<<DD MMM YYYY>>': pd.to_datetime(row['Compensation Effective Date']).strftime("%d %b %Y"),
#                 '<<FA>>': f"{format_number_indian(fixed)}",
#                 '<<VA>>': f"{format_number_indian(variable)}",
#                 '<<RA>>': f"{format_number_indian(retention)}" if retention else '',
#                 '<<TA>>': f"{format_number_indian(total_compensation)}",
#                 '<< INR Amount>>': f"INR {format_number_indian(int(str(row['ESPOS']).replace(',', '')))}" if row['ESPOS'] != 'No' else '',
#                 '<<Percentage>>': f"{int(row['Variable Pay - Payout']*100)}%",
#                 '<< Month>>': pd.to_datetime(row['Revised Pay effective month']).strftime("%B")
#             }


#             replace_placeholders_in_doc(doc, placeholders)

#             if row['Promotion'] != 'Yes':
#                 for paragraph in doc.paragraphs:
#                     if "You have been promoted as a" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             if row['Retention Pay'] == 'No':
#                 for paragraph in doc.paragraphs:
#                     if "Retention pay would be processed" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             if row['ESPOS'] == 'No':
#                 for paragraph in doc.paragraphs:
#                     if "You will be eligible for ESOPS worth" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             doc_start_time = time.time()
#             doc_path = os.path.join(temp_dir, f"{row['Name of Employee']}_Appraisal_letter.docx")
#             doc.save(doc_path)
#             doc_paths.append(doc_path)
#             doc_end_time = time.time()
#             print(f"Time taken to save word doc for {row['Name of Employee']}: {doc_end_time - doc_start_time}")
        
#         # Convert DOCX to PDF
#         pdf_conversion_start_time = time.time()
#         print(f"Start time for PDF Conversion: {pdf_conversion_start_time - start_time}")
#         for doc_path in doc_paths:
           
#             pdf_path = os.path.join(temp_dir + '_pdf', os.path.basename(doc_path).replace('.docx', '.pdf'))
#             os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
#             convert_docx_to_pdf(doc_path, pdf_path)
#             pdf_paths.append(pdf_path)
#         pdf_conversion_end_time = time.time()
#         print(f"Total time taken for PDF Conversion: {pdf_conversion_end_time - pdf_conversion_start_time}")

#         # Zip the files
#         zip_start_time = time.time()
#         st.session_state.doc_zip_path = os.path.join(temp_dir, "word_documents.zip")
#         st.session_state.pdf_zip_path = os.path.join(temp_dir, "pdf_documents.zip")
#         zip_files(doc_paths, st.session_state.doc_zip_path)
#         zip_files(pdf_paths, st.session_state.pdf_zip_path)
#         zip_end_time = time.time()
#         print(f"Time taken to zip the documents: {zip_end_time - zip_start_time}")
#         # Store paths in session state
#         st.session_state.doc_paths = doc_paths
#         st.session_state.pdf_paths = pdf_paths
#         total_time = time.time()
#         print(f"Total Time Taken : {total_time - start_time}")

#         st.success("Appraisal letters have been generated and are ready for download.")

# # Download buttons
# if st.session_state.doc_zip_path and os.path.exists(st.session_state.doc_zip_path):
#     with st.expander("Download All Word Documents"):
#         with open(st.session_state.doc_zip_path, "rb") as f:
#             st.download_button(
#                 label="Download All Appraisal Letters (Word)",
#                 data=f,
#                 file_name="Appraisal_word_documents.zip"
#             )

# if st.session_state.pdf_zip_path and os.path.exists(st.session_state.pdf_zip_path):
#     with st.expander("Download All PDF Documents"):
#         with open(st.session_state.pdf_zip_path, "rb") as f:
#             st.download_button(
#                 label="Download All Appraisal Letters (PDF)",
#                 data=f,
#                 file_name="Appraisal_pdf_documents.zip"
#             )

# if st.button("Clean up temporary files"):
#     st.session_state.cleanup = True
#     st.success("Temporary files will be cleaned up.")

# # Perform cleanup if flag is set
# if st.session_state.cleanup:
#     cleanup_temp_files("temp_files")
#     cleanup_temp_files("temp_files_pdf")
#     st.session_state.cleanup = False
#     st.session_state.doc_paths = []
#     st.session_state.pdf_paths = []
#     st.session_state.doc_zip_path = ""
#     st.session_state.pdf_zip_path = ""
#     st.success("Temporary files cleaned up.")

# import streamlit as st
# import pandas as pd
# from docx import Document
# import os
# from io import BytesIO
# # from docx2pdf import convert # Removed the docx2pdf import
# import pythoncom
# import zipfile
# import shutil
# import win32com.client
# import time

# st.title(":blue[Appraisal Letter] :orange[Generator]")

# # Function to convert DOCX to PDF using win32com
# def convert_docx_to_pdf(docx_file, pdf_path):
#     pythoncom.CoInitialize()
#     word = None  # Initialize word variable outside try block
#     doc = None # Initialize doc variable
#     try:
#         word = win32com.client.Dispatch("Word.Application")
#         word.Visible = False
#         doc = word.Documents.Open(docx_file)
#         doc.SaveAs(pdf_path, FileFormat=17) # 17 is for PDF
#     except Exception as e:
#         print(f"Error converting {docx_file} to PDF: {e}")
#     finally:
#         if doc:
#             doc.Close(False) # Close the document without saving any potential changes
#         if word:
#             word.Quit()   # Ensure word is closed properly.
#         pythoncom.CoUninitialize()
        
        

# # Function to replace placeholders in the document
# def replace_placeholders_in_doc(doc, placeholders):
#     # Replace in paragraphs
#     for paragraph in doc.paragraphs:
#         for placeholder, value in placeholders.items():
#             if placeholder in paragraph.text:
#                 paragraph.text = paragraph.text.replace(placeholder, str(value))
    
#     # Replace in table cells
#     for table in doc.tables:
#         rows_to_remove = []
#         for row_index, row in enumerate(table.rows):
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     for placeholder, value in placeholders.items():
#                         if placeholder in paragraph.text:
#                             paragraph.text = paragraph.text.replace(placeholder, str(value))
#                             if placeholder == '<<RA>>' and value == '':
#                                 rows_to_remove.append(row_index)
        
#         for row_index in reversed(rows_to_remove):
#             tbl = table._tbl
#             tbl.remove(tbl.tr_lst[row_index])

# # Function to format numbers in Indian numbering system
# def format_number_indian(number):
#     # Convert the number to a string
#     num_str = str(number)
    
#     # Split the number into integer and decimal parts
#     if '.' in num_str:
#         integer_part, decimal_part = num_str.split('.')
#     else:
#         integer_part, decimal_part = num_str, ''
    
#     # Reverse the integer part for easier processing
#     integer_part = integer_part[::-1]
    
#     # Insert commas according to the Indian numbering system
#     groups = []
#     groups.append(integer_part[:3])
#     integer_part = integer_part[3:]
    
#     while integer_part:
#         groups.append(integer_part[:2])
#         integer_part = integer_part[2:]
    
#     # Reverse the groups and join them with commas
#     formatted_integer = ','.join(groups)[::-1]
    
#     # Combine the integer and decimal parts
#     if decimal_part:
#         return formatted_integer + '.' + decimal_part
#     else:
#         return formatted_integer


# # Function to zip files
# def zip_files(file_paths, zip_path):
#     with zipfile.ZipFile(zip_path, 'w') as zipf:
#         for file in file_paths:
#             zipf.write(file, os.path.basename(file))

# # Function to clean up temporary files
# def cleanup_temp_files(temp_dir):
#     if os.path.exists(temp_dir):
#         # Remove all files in the directory
#         for file in os.listdir(temp_dir):
#             file_path = os.path.join(temp_dir, file)
#             if os.path.isfile(file_path):
#                 os.remove(file_path)
        
#         # Remove the directory itself
#         os.rmdir(temp_dir)

# # Initialize session state
# if 'doc_paths' not in st.session_state:
#     st.session_state.doc_paths = []
# if 'pdf_paths' not in st.session_state:
#     st.session_state.pdf_paths = []
# if 'doc_zip_path' not in st.session_state:
#     st.session_state.doc_zip_path = ""
# if 'pdf_zip_path' not in st.session_state:
#     st.session_state.pdf_zip_path = ""
# if 'cleanup' not in st.session_state:
#     st.session_state.cleanup = False

# # Upload the Excel file
# uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")

# if uploaded_file and not st.session_state.doc_paths:
#     # Define and create temporary directory
#     temp_dir = "temp_files"
#     os.makedirs(temp_dir, exist_ok=True)

#     doc_paths = []
#     pdf_paths = []
#     start_time = time.time()
#     print(f"Start Time: {start_time}")


#     with st.spinner('Processing the Excel file...'):
#         df = pd.read_excel(uploaded_file)
#         template_path = 'Compensation Revision Letter_Format.docx'
#         excel_read_time = time.time()
#         print(f"Time to read excel: {excel_read_time - start_time}")

#         for index, row in df.iterrows():
#             doc = Document(template_path)

#             fixed = int(str(row['Fixed Pay']).replace(',', ''))
#             variable = int(str(row['Variable Pay']).replace(',', ''))
#             retention = 0 if row['Retention Pay'] == 'No' else int(str(row['Retention Pay']).replace(',', ''))
#             total_compensation = fixed + variable + retention

#             placeholders = {
#                 '<<Month YYYY>>': pd.to_datetime(row['Month of the Letter issued']).strftime("%d %b %Y"),
#                 '<<Name>>': row['Name of Employee'],
#                 '<<Designation>>': row['New Designation'] if row['Promotion'] == 'Yes' else '',
#                 '<<DD MMM YYYY>>': pd.to_datetime(row['Compensation Effective Date']).strftime("%d %b %Y"),
#                 '<<FA>>': f"{format_number_indian(fixed)}",
#                 '<<VA>>': f"{format_number_indian(variable)}",
#                 '<<RA>>': f"{format_number_indian(retention)}" if retention else '',
#                 '<<TA>>': f"{format_number_indian(total_compensation)}",
#                 '<< INR Amount>>': f"INR {format_number_indian(int(str(row['ESPOS']).replace(',', '')))}" if row['ESPOS'] != 'No' else '',
#                 '<<Percentage>>': f"{int(row['Variable Pay - Payout']*100)}%",
#                 '<< Month>>': pd.to_datetime(row['Revised Pay effective month']).strftime("%B")
#             }


#             replace_placeholders_in_doc(doc, placeholders)

#             if row['Promotion'] != 'Yes':
#                 for paragraph in doc.paragraphs:
#                     if "You have been promoted as a" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             if row['Retention Pay'] == 'No':
#                 for paragraph in doc.paragraphs:
#                     if "Retention pay would be processed" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             if row['ESPOS'] == 'No':
#                 for paragraph in doc.paragraphs:
#                     if "You will be eligible for ESOPS worth" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             doc_start_time = time.time()
#             doc_path = os.path.join(temp_dir, f"{row['Name of Employee']}_Appraisal_letter.docx")
#             doc.save(doc_path)
#             # time.sleep(0.1)
#             doc_paths.append(doc_path)
#             doc_end_time = time.time()
#             print(f"Time taken to save word doc for {row['Name of Employee']}: {doc_end_time - doc_start_time}")
        
#         # Convert DOCX to PDF
#         pdf_conversion_start_time = time.time()
#         print(f"Start time for PDF Conversion: {pdf_conversion_start_time - start_time}")
#         for doc_path in doc_paths:
           
#             pdf_path = os.path.join(temp_dir + '_pdf', os.path.basename(doc_path).replace('.docx', '.pdf'))
#             os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
#             convert_docx_to_pdf(doc_path, pdf_path)
#             pdf_paths.append(pdf_path)
#         pdf_conversion_end_time = time.time()
#         print(f"Total time taken for PDF Conversion: {pdf_conversion_end_time - pdf_conversion_start_time}")

#         # Zip the files
#         zip_start_time = time.time()
#         st.session_state.doc_zip_path = os.path.join(temp_dir, "word_documents.zip")
#         st.session_state.pdf_zip_path = os.path.join(temp_dir, "pdf_documents.zip")
#         zip_files(doc_paths, st.session_state.doc_zip_path)
#         zip_files(pdf_paths, st.session_state.pdf_zip_path)
#         zip_end_time = time.time()
#         print(f"Time taken to zip the documents: {zip_end_time - zip_start_time}")
#         # Store paths in session state
#         st.session_state.doc_paths = doc_paths
#         st.session_state.pdf_paths = pdf_paths
#         total_time = time.time()
#         print(f"Total Time Taken : {total_time - start_time}")

#         st.success("Appraisal letters have been generated and are ready for download.")

# # Download buttons
# if st.session_state.doc_zip_path and os.path.exists(st.session_state.doc_zip_path):
#     with st.expander("Download All Word Documents"):
#         with open(st.session_state.doc_zip_path, "rb") as f:
#             st.download_button(
#                 label="Download All Appraisal Letters (Word)",
#                 data=f,
#                 file_name="Appraisal_word_documents.zip"
#             )

# if st.session_state.pdf_zip_path and os.path.exists(st.session_state.pdf_zip_path):
#     with st.expander("Download All PDF Documents"):
#         with open(st.session_state.pdf_zip_path, "rb") as f:
#             st.download_button(
#                 label="Download All Appraisal Letters (PDF)",
#                 data=f,
#                 file_name="Appraisal_pdf_documents.zip"
#             )

# if st.button("Clean up temporary files"):
#     st.session_state.cleanup = True
#     st.success("Temporary files will be cleaned up.")

# # Perform cleanup if flag is set
# if st.session_state.cleanup:
#     cleanup_temp_files("temp_files")
#     cleanup_temp_files("temp_files_pdf")
#     st.session_state.cleanup = False
#     st.session_state.doc_paths = []
#     st.session_state.pdf_paths = []
#     st.session_state.doc_zip_path = ""
#     st.session_state.pdf_zip_path = ""
#     st.success("Temporary files cleaned up.")


# import streamlit as st
# import pandas as pd
# from docx import Document
# import os
# from io import BytesIO
# import pythoncom
# import zipfile
# import shutil
# import win32com.client
# import time

# st.title("Appraisal Letter Generator")

# # Function to convert DOCX to PDF using win32com
# def convert_docx_to_pdf(docx_file, pdf_path):
#     pythoncom.CoInitialize()
#     word = None
#     doc = None
#     try:
#         word = win32com.client.Dispatch("Word.Application")
#         word.Visible = False
#         doc = word.Documents.Open(docx_file)
#         doc.SaveAs(pdf_path, FileFormat=17) # 17 is for PDF
#     except Exception as e:
#         print(f"Error converting {docx_file} to PDF: {e}")
#     finally:
#         if doc:
#             doc.Close(False)
#         if word:
#             word.Quit()
#         pythoncom.CoUninitialize()
        
        

# # Function to replace placeholders in the document
# def replace_placeholders_in_doc(doc, placeholders):
#     # Replace in paragraphs
#     for paragraph in doc.paragraphs:
#         for placeholder, value in placeholders.items():
#             if placeholder in paragraph.text:
#                 paragraph.text = paragraph.text.replace(placeholder, str(value))
    
#     # Replace in table cells
#     for table in doc.tables:
#         rows_to_remove = []
#         for row_index, row in enumerate(table.rows):
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     for placeholder, value in placeholders.items():
#                         if placeholder in paragraph.text:
#                             paragraph.text = paragraph.text.replace(placeholder, str(value))
#                             if placeholder == '<<RA>>' and value == '':
#                                 rows_to_remove.append(row_index)
        
#         for row_index in reversed(rows_to_remove):
#             tbl = table._tbl
#             tbl.remove(tbl.tr_lst[row_index])

# # Function to format numbers in Indian numbering system
# def format_number_indian(number):
#     # Convert the number to a string
#     num_str = str(number)
    
#     # Split the number into integer and decimal parts
#     if '.' in num_str:
#         integer_part, decimal_part = num_str.split('.')
#     else:
#         integer_part, decimal_part = num_str, ''
    
#     # Reverse the integer part for easier processing
#     integer_part = integer_part[::-1]
    
#     # Insert commas according to the Indian numbering system
#     groups = []
#     groups.append(integer_part[:3])
#     integer_part = integer_part[3:]
    
#     while integer_part:
#         groups.append(integer_part[:2])
#         integer_part = integer_part[2:]
    
#     # Reverse the groups and join them with commas
#     formatted_integer = ','.join(groups)[::-1]
    
#     # Combine the integer and decimal parts
#     if decimal_part:
#         return formatted_integer + '.' + decimal_part
#     else:
#         return formatted_integer


# # Function to zip files
# def zip_files(file_paths, zip_path):
#     with zipfile.ZipFile(zip_path, 'w') as zipf:
#         for file in file_paths:
#             zipf.write(file, os.path.basename(file))

# # Function to clean up temporary files
# def cleanup_temp_files(temp_dir):
#     if os.path.exists(temp_dir):
#         # Remove all files in the directory
#         for file in os.listdir(temp_dir):
#             file_path = os.path.join(temp_dir, file)
#             if os.path.isfile(file_path):
#                 os.remove(file_path)
        
#         # Remove the directory itself
#         os.rmdir(temp_dir)

# # Initialize session state
# if 'doc_paths' not in st.session_state:
#     st.session_state.doc_paths = []
# if 'pdf_paths' not in st.session_state:
#     st.session_state.pdf_paths = []
# if 'doc_zip_path' not in st.session_state:
#     st.session_state.doc_zip_path = ""
# if 'pdf_zip_path' not in st.session_state:
#     st.session_state.pdf_zip_path = ""
# if 'cleanup' not in st.session_state:
#     st.session_state.cleanup = False

# # Upload the Excel file
# uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")

# if uploaded_file and not st.session_state.doc_paths:
#     # Define and create temporary directory
#     temp_dir = os.path.join(os.getcwd(), "temp_files")
#     os.makedirs(temp_dir, exist_ok=True)

#     temp_pdf_dir = os.path.join(os.getcwd(), "temp_files_pdf")
#     os.makedirs(temp_pdf_dir, exist_ok=True)

#     doc_paths = []
#     pdf_paths = []
#     start_time = time.time()
#     print(f"Start Time: {start_time}")


#     with st.spinner('Processing the Excel file...'):
#         df = pd.read_excel(uploaded_file)
#         template_path = 'Compensation Revision Letter_Format.docx'
#         excel_read_time = time.time()
#         print(f"Time to read excel: {excel_read_time - start_time}")

#         for index, row in df.iterrows():
#             doc = Document(template_path)

#             fixed = int(str(row['Fixed Pay']).replace(',', ''))
#             variable = int(str(row['Variable Pay']).replace(',', ''))
#             retention = 0 if row['Retention Pay'] == 'No' else int(str(row['Retention Pay']).replace(',', ''))
#             total_compensation = fixed + variable + retention

#             placeholders = {
#                 '<<Month YYYY>>': pd.to_datetime(row['Month of the Letter issued']).strftime("%d %b %Y"),
#                 '<<Name>>': row['Name of Employee'],
#                 '<<Designation>>': row['New Designation'] if row['Promotion'] == 'Yes' else '',
#                 '<<DD MMM YYYY>>': pd.to_datetime(row['Compensation Effective Date']).strftime("%d %b %Y"),
#                 '<<FA>>': f"{format_number_indian(fixed)}",
#                 '<<VA>>': f"{format_number_indian(variable)}",
#                 '<<RA>>': f"{format_number_indian(retention)}" if retention else '',
#                 '<<TA>>': f"{format_number_indian(total_compensation)}",
#                 '<< INR Amount>>': f"INR {format_number_indian(int(str(row['ESPOS']).replace(',', '')))}" if row['ESPOS'] != 'No' else '',
#                 '<<Percentage>>': f"{int(row['Variable Pay - Payout']*100)}%",
#                 '<< Month>>': pd.to_datetime(row['Revised Pay effective month']).strftime("%B")
#             }


#             replace_placeholders_in_doc(doc, placeholders)

#             if row['Promotion'] != 'Yes':
#                 for paragraph in doc.paragraphs:
#                     if "You have been promoted as a" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             if row['Retention Pay'] == 'No':
#                 for paragraph in doc.paragraphs:
#                     if "Retention pay would be processed" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             if row['ESPOS'] == 'No':
#                 for paragraph in doc.paragraphs:
#                     if "You will be eligible for ESOPS worth" in paragraph.text:
#                         p = paragraph._element
#                         p.getparent().remove(p)

#             doc_start_time = time.time()
#             doc_path = os.path.join(temp_dir, f"{row['Name of Employee']}_Appraisal_letter.docx")
#             doc.save(doc_path)
#             time.sleep(0.1)
#             doc_paths.append(doc_path)
#             doc_end_time = time.time()
#             print(f"Time taken to save word doc for {row['Name of Employee']}: {doc_end_time - doc_start_time}")
        
#         # Convert DOCX to PDF
#         pdf_conversion_start_time = time.time()
#         print(f"Start time for PDF Conversion: {pdf_conversion_start_time - start_time}")
#         for doc_path in doc_paths:
#             pdf_path = os.path.join(os.getcwd(), temp_dir + '_pdf', os.path.basename(doc_path).replace('.docx', '.pdf'))
#             os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
            
#             retries = 3
#             converted = False
#             while retries > 0:
#                 if os.path.exists(doc_path):
#                   try:
#                       convert_docx_to_pdf(doc_path, pdf_path)
#                       converted=True
#                       break #success, move to next document
#                   except Exception as e:
#                       print(f"Error converting {doc_path} on attempt {3 - retries}: {e}")
#                       retries -=1
#                       time.sleep(0.1) # Small delay before retry
#                 else:
#                    print(f"File {doc_path} not found, retrying")
#                    retries-=1
#                    time.sleep(0.1) # Small delay before retry
#             else:
#                 print(f"Failed to convert {doc_path} after multiple retries")
            
#             if converted:
#                 pdf_paths.append(pdf_path)
            
#         pdf_conversion_end_time = time.time()
#         print(f"Total time taken for PDF Conversion: {pdf_conversion_end_time - pdf_conversion_start_time}")

#         # Zip the files
#         zip_start_time = time.time()
#         st.session_state.doc_zip_path = os.path.join(temp_dir, "word_documents.zip")
#         st.session_state.pdf_zip_path = os.path.join(temp_pdf_dir, "pdf_documents.zip")
#         zip_files(doc_paths, st.session_state.doc_zip_path)
#         zip_files(pdf_paths, st.session_state.pdf_zip_path)
#         zip_end_time = time.time()
#         print(f"Time taken to zip the documents: {zip_end_time - zip_start_time}")
#         # Store paths in session state
#         st.session_state.doc_paths = doc_paths
#         st.session_state.pdf_paths = pdf_paths
#         total_time = time.time()
#         print(f"Total Time Taken : {total_time - start_time}")

#         st.success("Appraisal letters have been generated and are ready for download.")

# # Download buttons
# if st.session_state.doc_zip_path and os.path.exists(st.session_state.doc_zip_path):
#     with st.expander("Download All Word Documents"):
#         with open(st.session_state.doc_zip_path, "rb") as f:
#             st.download_button(
#                 label="Download All Appraisal Letters (Word)",
#                 data=f,
#                 file_name="Appraisal_word_documents.zip"
#             )

# if st.session_state.pdf_zip_path and os.path.exists(st.session_state.pdf_zip_path):
#     with st.expander("Download All PDF Documents"):
#         with open(st.session_state.pdf_zip_path, "rb") as f:
#             st.download_button(
#                 label="Download All Appraisal Letters (PDF)",
#                 data=f,
#                 file_name="Appraisal_pdf_documents.zip"
#             )

# if st.button("Clean up temporary files"):
#     st.session_state.cleanup = True
#     st.success("Temporary files will be cleaned up.")

# # Perform cleanup if flag is set
# if st.session_state.cleanup:
#     cleanup_temp_files("temp_files")
#     cleanup_temp_files("temp_files_pdf")
#     st.session_state.cleanup = False
#     st.session_state.doc_paths = []
#     st.session_state.pdf_paths = []
#     st.session_state.doc_zip_path = ""
#     st.session_state.pdf_zip_path = ""
#     st.success("Temporary files cleaned up.")

import streamlit as st
import pandas as pd
from docx import Document
import os
from io import BytesIO
import pythoncom
import zipfile
import shutil
import win32com.client
import time

st.title("Appraisal Letter Generator")

# Function to convert DOCX to PDF using win32com
def convert_docx_to_pdf(docx_file, pdf_path):
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_file)
        doc.SaveAs(pdf_path, FileFormat=17) # 17 is for PDF
    except Exception as e:
        print(f"Error converting {docx_file} to PDF: {e}")
    finally:
        if doc:
            doc.Close(False)
        if word:
            word.Quit()
        pythoncom.CoUninitialize()
        
        

# Function to replace placeholders in the document
def replace_placeholders_in_doc(doc, placeholders):
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
    
    # Replace in table cells
    for table in doc.tables:
        rows_to_remove = []
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholders.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
                            if placeholder == '<<RA>>' and value == '':
                                rows_to_remove.append(row_index)
        
        for row_index in reversed(rows_to_remove):
            tbl = table._tbl
            tbl.remove(tbl.tr_lst[row_index])

# Function to format numbers in Indian numbering system
def format_number_indian(number):
    # Convert the number to a string
    num_str = str(number)
    
    # Split the number into integer and decimal parts
    if '.' in num_str:
        integer_part, decimal_part = num_str.split('.')
    else:
        integer_part, decimal_part = num_str, ''
    
    # Reverse the integer part for easier processing
    integer_part = integer_part[::-1]
    
    # Insert commas according to the Indian numbering system
    groups = []
    groups.append(integer_part[:3])
    integer_part = integer_part[3:]
    
    while integer_part:
        groups.append(integer_part[:2])
        integer_part = integer_part[2:]
    
    # Reverse the groups and join them with commas
    formatted_integer = ','.join(groups)[::-1]
    
    # Combine the integer and decimal parts
    if decimal_part:
        return formatted_integer + '.' + decimal_part
    else:
        return formatted_integer


# Function to zip files
def zip_files(file_paths, zip_path):
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in file_paths:
            zipf.write(file, os.path.basename(file))

# Function to clean up temporary files
def cleanup_temp_files(temp_dir):
    if os.path.exists(temp_dir):
        # Remove all files in the directory
        for file in os.listdir(temp_dir):
            file_path = os.path.join(temp_dir, file)
            if os.path.isfile(file_path):
                os.remove(file_path)
        
        # Remove the directory itself
        os.rmdir(temp_dir)

# Initialize session state
if 'doc_paths' not in st.session_state:
    st.session_state.doc_paths = []
if 'pdf_paths' not in st.session_state:
    st.session_state.pdf_paths = []
if 'doc_zip_path' not in st.session_state:
    st.session_state.doc_zip_path = ""
if 'pdf_zip_path' not in st.session_state:
    st.session_state.pdf_zip_path = ""
if 'cleanup' not in st.session_state:
    st.session_state.cleanup = False

# Upload the Excel file
uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")

if uploaded_file and not st.session_state.doc_paths:
    # Define and create temporary directory
    temp_dir = os.path.join(os.getcwd(), "temp_files")
    os.makedirs(temp_dir, exist_ok=True)

    temp_pdf_dir = os.path.join(os.getcwd(), "temp_files_pdf")
    os.makedirs(temp_pdf_dir, exist_ok=True)

    doc_paths = []
    pdf_paths = []
    start_time = time.time()
    print(f"Start Time: {start_time}")


    with st.spinner('Processing the Excel file...'):
        df = pd.read_excel(uploaded_file)
        template_path = 'Compensation Revision Letter_Format.docx'
        excel_read_time = time.time()
        print(f"Time to read excel: {excel_read_time - start_time}")

        for index, row in df.iterrows():
            doc = Document(template_path)

            fixed = int(str(row['Fixed Pay']).replace(',', ''))
            variable = int(str(row['Variable Pay']).replace(',', ''))
            retention = 0 if row['Retention Pay'] == 'No' else int(str(row['Retention Pay']).replace(',', ''))
            total_compensation = fixed + variable + retention

            placeholders = {
                '<<Month YYYY>>': pd.to_datetime(row['Month of the Letter issued']).strftime("%d %b %Y"),
                '<<Name>>': row['Name of Employee'],
                '<<Designation>>': row['New Designation'] if row['Promotion'] == 'Yes' else '',
                '<<DD MMM YYYY>>': pd.to_datetime(row['Compensation Effective Date']).strftime("%d %b %Y"),
                '<<FA>>': f"{format_number_indian(fixed)}",
                '<<VA>>': f"{format_number_indian(variable)}",
                '<<RA>>': f"{format_number_indian(retention)}" if retention else '',
                '<<TA>>': f"{format_number_indian(total_compensation)}",
                '<< INR Amount>>': f"INR {format_number_indian(int(str(row['ESPOS']).replace(',', '')))}" if row['ESPOS'] != 'No' else '',
                '<<Percentage>>': f"{int(row['Variable Pay - Payout']*100)}%",
                '<< Month>>': pd.to_datetime(row['Revised Pay effective month']).strftime("%B")
            }


            replace_placeholders_in_doc(doc, placeholders)

            if row['Promotion'] != 'Yes':
                for paragraph in doc.paragraphs:
                    if "You have been promoted as a" in paragraph.text:
                        p = paragraph._element
                        p.getparent().remove(p)

            if row['Retention Pay'] == 'No':
                for paragraph in doc.paragraphs:
                    if "Retention pay would be processed" in paragraph.text:
                        p = paragraph._element
                        p.getparent().remove(p)

            if row['ESPOS'] == 'No':
                for paragraph in doc.paragraphs:
                    if "You will be eligible for ESOPS worth" in paragraph.text:
                        p = paragraph._element
                        p.getparent().remove(p)

            doc_start_time = time.time()
            doc_path = os.path.join(temp_dir, f"{row['Name of Employee']}_Appraisal_letter.docx")
            doc.save(doc_path)
            doc_paths.append(doc_path)
            doc_end_time = time.time()
            print(f"Time taken to save word doc for {row['Name of Employee']}: {doc_end_time - doc_start_time}")
        
        # Convert DOCX to PDF
        pdf_conversion_start_time = time.time()
        print(f"Start time for PDF Conversion: {pdf_conversion_start_time - start_time}")
        for doc_path in doc_paths:
            pdf_path = os.path.join(os.getcwd(), temp_dir + '_pdf', os.path.basename(doc_path).replace('.docx', '.pdf'))
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
            convert_docx_to_pdf(doc_path, pdf_path)
            pdf_paths.append(pdf_path)
        pdf_conversion_end_time = time.time()
        print(f"Total time taken for PDF Conversion: {pdf_conversion_end_time - pdf_conversion_start_time}")

        # Zip the files
        zip_start_time = time.time()
        st.session_state.doc_zip_path = os.path.join(temp_dir, "word_documents.zip")
        st.session_state.pdf_zip_path = os.path.join(temp_pdf_dir, "pdf_documents.zip")
        zip_files(doc_paths, st.session_state.doc_zip_path)
        zip_files(pdf_paths, st.session_state.pdf_zip_path)
        zip_end_time = time.time()
        print(f"Time taken to zip the documents: {zip_end_time - zip_start_time}")
        # Store paths in session state
        st.session_state.doc_paths = doc_paths
        st.session_state.pdf_paths = pdf_paths
        total_time = time.time()
        print(f"Total Time Taken : {total_time - start_time}")

        st.success("Appraisal letters have been generated and are ready for download.")

# Download buttons
if st.session_state.doc_zip_path and os.path.exists(st.session_state.doc_zip_path):
    with st.expander("Download All Word Documents"):
        with open(st.session_state.doc_zip_path, "rb") as f:
            st.download_button(
                label="Download All Appraisal Letters (Word)",
                data=f,
                file_name="Appraisal_word_documents.zip"
            )

if st.session_state.pdf_zip_path and os.path.exists(st.session_state.pdf_zip_path):
    with st.expander("Download All PDF Documents"):
        with open(st.session_state.pdf_zip_path, "rb") as f:
            st.download_button(
                label="Download All Appraisal Letters (PDF)",
                data=f,
                file_name="Appraisal_pdf_documents.zip"
            )

if st.button("Clean up temporary files"):
    st.session_state.cleanup = True
    st.success("Temporary files will be cleaned up.")

# Perform cleanup if flag is set
if st.session_state.cleanup:
    cleanup_temp_files("temp_files")
    cleanup_temp_files("temp_files_pdf")
    st.session_state.cleanup = False
    st.session_state.doc_paths = []
    st.session_state.pdf_paths = []
    st.session_state.doc_zip_path = ""
    st.session_state.pdf_zip_path = ""
    st.success("Temporary files cleaned up.")